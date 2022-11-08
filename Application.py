# Sources: Sweigart, Al. Automate the Boring Stuff with Python: Practical Programming for Total Beginners. No Starch Press, 2015. 

import docx
from docx.shared import Pt
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import sys
import os

class WordDocument:
    def changeDate(self, newDate, filename):
        doc = docx.Document(filename)
        dateParagraph = doc.paragraphs[8].text # the paragraph where the date is
        indexOfColon = dateParagraph.find(":") # gets the index where I begin to write in the date

        date_substring = dateParagraph[indexOfColon+2:] # the old date substring
        newString = dateParagraph.replace(date_substring, newDate) # stores a string that replaces the old date with the new date
        doc.paragraphs[8].text = newString # In the doc, replace the old string with the new string that has the new date
        doc.save(filename)

    def changeStyle(self, filename, paragraphNum):
        doc = docx.Document(filename)
        run = doc.paragraphs[8].add_run()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial Black"
        font.size = Pt(10)
        doc.save(filename)

class Email:
    my_email_address = ""
    recipient_email_address = ""
    password = ""
    subject = ""
    email_body = ""
    msg = MIMEMultipart()

    def __init__(self, my_email_address, password, recipient_email_address, subject, email_body):
        self.my_email_address = my_email_address
        self.recipient_email_address = recipient_email_address
        self.password = password
        self.subject = subject
        self.email_body = email_body

    def insert_attachment(self, filename):
        self.msg["From"] = self.my_email_address
        self.msg["To"] = self.recipient_email_address
        self.msg["Subject"] = self.subject

        self.msg.attach(MIMEText(self.email_body, "plain"))
        attachment = open(filename, "rb")

        p = MIMEBase("application", "octet-stream")
        p.set_payload(attachment.read())

        encoders.encode_base64(p)
        p.add_header("Content-Disposition", "attachment; filename= %s" % filename)

        self.msg.attach(p)

    def send(self):
        smtpObj = smtplib.SMTP("smtp.gmail.com", 587) # Email protocol
        smtpObj.starttls() # Enable encryption
        smtpObj.login(self.my_email_address, self.password)
        sent_mail = False

        text = self.msg.as_string()
        dict = smtpObj.sendmail(self.my_email_address, self.recipient_email_address, text)

        # If dict is empty, then sending the e-mail was successful
        if len(dict) == 0:
            sent_mail = True

        smtpObj.quit() # Disconnect from e-mail server

        return sent_mail

    # Getters and setters
    def setMyEmailAddress(self, my_email_address):
        self.my_email_address = my_email_address

    def setRecipientEmailAddress(self, recipient_email_address):
        self.recipient_email_address = recipient_email_address

    def setPassword(self, password):
        self.password = password

    def setSubject(self, subject):
        self.subject = subject

    def setEmailBody(self, email_body):
        self.email_body = email_body

    def getMyEmailAddress(self):
        return self.my_email_address

    def getRecipientEmailAddress(self):
        return self.recipient_email_address

    def getSubject(self):
        return self.subject

    def getEmailBody(self):
        return self.email_body

    def getPassword(self):
        return self.password

my_email_address = "chaseconner366@gmail.com"
password = input("Please enter your password: ")
selection = input("\nSend to:\n\n   1) Metro Mobility\n   2) Another recipient\n\nEnter here: ")
recipient_email_address = ""

if selection == "1":
    recipient_email_address = "WReservations@metc.state.mn.us"
elif selection == "2":
     print()
     recipient_email_address = input("Please enter recipient email address: ")

subject = "Email Reservation Forms Attached"
email_body = "Dear Metro Mobility,\n\nI have attached my e-mail reservation form(s)" \
             " to this e-mail. You may call me at 763-202-1523 to let me" \
             " know the pickup time(s) the day before. You may also give my " \
             "number to the bus driver in case he/she has trouble locating " \
             "me.\n\nThank you,\nChase Conner"

email = Email(my_email_address, password, recipient_email_address, subject, email_body)

attachAnotherForm = True

while attachAnotherForm == True:
    doc = WordDocument()
    listDir = os.listdir() # List of files in current working directory
    count = 0 # Prints a file number next to the file name
    print() # Put a space 
    for file in listDir:
        print(str(count) + ") " + file)
        count+=1 

    selection = input("\nEnter file number here: ")
    filename = listDir[int(selection)]
    newDate = input("For " + filename + " enter new date like this - Thursday, 10/2/18. Enter here: ")
    doc.changeDate(newDate, filename)

    email.insert_attachment(filename)

    selection = input("Attach another form to email? Y/N or y/n: ")
    selection = selection.upper()
    if selection == "Y":
        attachAnotherForm = True
        email.setEmailBody("")
        email.setSubject("")
    elif selection == "N":
        attachAnotherForm = False
    else:
        sys.exit("Invalid input.")

sent_mail = email.send()

if sent_mail == True:
    print("\nE-mail to " + email.getRecipientEmailAddress() + " was successfully sent.")
else:
    print("E-mail " + email.getRecipientEmailAddress() + " failed to send.")

input("\nPress [ENTER] to exit")
sys.exit()